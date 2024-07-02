from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from CR_add_funcs import add_contact_info, add_education, add_experience, add_projects, add_sectionheader, add_technical_skills

para_spacing = 6
line_spacing = 14



def create_resume(input):
    document = Document()

    # Set margins
    left_marg = right_marg = 0.5
    document.sections[0].top_margin = Cm(1.11)
    document.sections[0].bottom_margin = Cm(0.73)
    document.sections[0].left_margin = Inches(left_marg)
    document.sections[0].right_margin = Inches(right_marg)
    
    # Set font size, font, and alignment for the name
    name_paragraph = document.add_paragraph()
    name_run = name_paragraph.add_run(input[0][0])
    name_run.bold = True
    name_run.font.size = Pt(input[0][2])
    name_run.font.name = input[0][1]
    name_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    name_paragraph.paragraph_format.space_after = Pt(0)  # Remove space after paragraph
    

    add_contact_info(document, input[1])
  
    for sections in input[2]:
        print(sections[2][0][0])
        if ((sections[2][0][0]) == ''):
            continue
        add_sectionheader(document, sections[0][0], sections[0][1], sections[0][2])
        for parts in sections[2]:
            if sections[0][0] == "EDUCATION":
                sections[1](document, parts[0], parts[1], parts[2], parts[3], parts[4], sections[0][1], sections[0][2])
            if sections[0][0] == "EXPERIENCE":
                sections[1](document, parts[0], parts[1], parts[2], parts[3], parts[4], sections[0][1], sections[0][2])
            if sections[0][0] == "PROJECTS":
                sections[1](document, parts[0], parts[1], parts[2], sections[0][1], sections[0][2])
            if sections[0][0] == "TECHNICAL SKILLS":
                 sections[1](document, parts, sections[0][1], sections[0][2])
    
    document.save("JakeCV.docx")
    #pypandoc.convert_file('JakeCV.docx', 'pdf', outputfile='templates/JakeCV.pdf')
    print("Resume created successfully!")




resume_array = [["Aryan Patel", "Garamond", 24], 
                ["aryan.patel2@uwaterloo.ca", "(437) 339-5154", "github.com/aryanp05", "linkedin.com/in/aryanpatel05"],
                [
                    [
                        ["EDUCATION", "Garamond", [12, 11, 10]],
                        add_education,
                        [
                            ["The University of Waterloo", "Waterloo, ON", 
                            "Bachelors of Computer Science and Bachelors of Business Administration with 4.0 GPA", "April 2028", 
                            ["Relevant Coursework: Function Programing, Tools for Software Dev, Algorithm Design & Data Abstraction",
                            "Extracurricular: Hackathon Executive, Computer Science Club, Data Science Club"]]
                        ]
                    ],
                    [
                        ["EXPERIENCE", "Garamond", [12, 11, 11]],
                        add_experience,
                        [
                            ["Software Developer & LLM Engineer", "Sep. 2023 - March 2024", "A.P.U.G. Club", 
                             "Waterloo, ON", [
                            "Developed LooLearnAI  (), a student-focus AI software that assists in academics by answering content-specific questions using course conventions and generating practise material using **OpenAI API, Milvus, and Python**",
                            "Implemented sophisticated Natural Language Processing (NLP) techniques for semantic search and data vectorization, ensuring intuitive query understanding and resource discovery, to increase retrieval speed by 25%.",
                            "Employed streamlined deployment and consistent application performance using the tool Docker.",
                            "Used advanced Prompt Engineering to refine the AI insights into user queries and generating solutions and developed target algorithms including **R.A.G. & R.S.E.** to increase LooLearnAI answer accuracy from 78% to 93%"]],
                            
                            ["Software Developer", "Sep. 2020 - Nov. 2021", "Neonic Wraps", "Toronto, Ontario",
                            [
                            "Designed and developed a fully functional and attractive e-commerce website using HTML, CSS, and JavaScript to increase customer conversion by 80%",
                            "Managed online inventory and performed database queries using SQL on MongoDB to effectively track items and orders",
                            "Operated an e-commerce business selling a diverse range of tech products, achieving $10k+ in sales."]]
                        
                        ]
                    ],
                    [
                        ["PROJECTS", "Garamond", [12, 11, 10]],
                        add_projects,
                        [
                            ["Wildfire Path Predictor", "May 2024 - June 2024", [
                            "Engineered the Wildfire Path Predictor for the Wildfire AI Hackathon 2024, utilizing PyTorch and Torchvision to create a hybrid model combining CNNs and LSTMs, accurately predicting wildfire spread by processing spatial and temporal data from satellite imagery and climate variables.",
                            "Leveraged satellite data from MODIS, VIIRS, and IBAND to train AI models, incorporating fire weather data, enhancing wildfire prediction accuracy."]],
                            
                            ["GetTrash Garbage Identifier", "May 2024 - May 2024", [
                            "Created GetTrash, a web app leveraging Python, TensorFlow, and Streamlit for real-time waste sorting through live video scanning and image captioning, achieving third place at the Google Developer SC-hosted Hack With AI event.",
                            "Created a robust garbage classification model using TensorFlow, trained a CNN architecture and leveraging large online data sets of over 16,000+ photos to identify trash to a high accuracy"]],
                            
                            ["Meerkat ASL", "April 2024 - May 2024", [
                            "Developed a sign language interpreter with machine learning techniques and Scikit-Learn leveraging OpenCV to capture a dataset of over 40,000 images",
                            "Implemented a preprocessing pipeline to extract and preprocess hand landmarks using Mediapipe achieving efficient representation of sign language gestures for input into a neural network model"]]
                        ]
                    ],
                    [
                        ["TECHNICAL SKILLS", "Garamond", [12, 11, 10]],
                        add_technical_skills,
                        [
                            ["Programming Languages", "Python, C++, C, HTML, CSS, JavaScript, Bash, Git, Racket, SQL"],
                            ["Technologies", "Flask, OpenCV, Tensorflow, PyTorch, Docker, Kubernetes, Git, Milvus, Rasterio, Streamlit"]
                        ]
                    ]

                ]
            ]

#create_resume(resume_array)