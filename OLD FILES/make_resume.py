from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from CR_spacing import calculate_num_of_space
#from add_part import add_contact_info, add_education, add_experience, add_projects, add_sectionheader, add_technical_skills

para_spacing = 6
line_spacing = 14

def create_resume():
    document = Document()

    # Set margins
    left_marg = right_marg = 0.5
    document.sections[0].top_margin = Cm(1.11)
    document.sections[0].bottom_margin = Cm(0.73)
    document.sections[0].left_margin = Inches(left_marg)
    document.sections[0].right_margin = Inches(right_marg)
    
    # Set font size, font, and alignment for the name
    name_paragraph = document.add_paragraph()
    name_run = name_paragraph.add_run("Jake Ryan")
    name_run.bold = True
    name_run.font.size = Pt(24)
    name_run.font.name = 'Garamond'
    name_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    name_paragraph.paragraph_format.space_after = Pt(0)  # Remove space after paragraph
    
    # Add contact information
    #add_contact_info(document, "(123) 456-7890", "jake@su.edu", "linkedin.com/in/jake", "github.com/jake")
    add_contact_info(document, "jake@su.edu", "(123) 456-7890", "github.com/jake", "linkedin.com/in/jake")
    # Add education section
    add_sectionheader(document, "EDUCATION")
    add_education(document, "The University of Waterloo", "Waterloo, ON", 
                  "Bachelors of Computer Science and Bachelors of Business Administration with 4.0 GPA", "April 2028", 
                  ["Relevant Coursework: Function Programing, Tools for Software Dev, Algorithm Design & Data Abstraction",
                   "Extracurricular: Hackathon Executive, Computer Science Club, Data Science Club"])
    #document.add_paragraph()
    
    # Add experience section
    add_sectionheader(document, "EXPERIENCE")
    add_experience(document, "Undergraduate Research Assistant", "Texas A&M University", "College Station, TX",
                   "June 2020 – Present", [
                       "Developed a REST API using FastAPI and PostgreSQL to store data from learning management systems",
                       "Developed a full-stack web application using Flask, React, PostgreSQL and Docker to analyze GitHub data",
                       "Explored ways to visualize GitHub collaboration in a classroom setting"
                   ])
    add_experience(document, "Information Technology Support Specialist", "Southwestern University", "Georgetown, TX",
                   "Sep. 2018 – Present", [
                       "Communicate with managers to set up campus computers used on campus",
                       "Assess and troubleshoot computer problems brought by students, faculty and staff",
                       "Maintain upkeep of computers, classroom equipment, and 200 printers across campus"
                   ])
    add_experience(document, "Artificial Intelligence Research Assistant", "Southwestern University", "Georgetown, TX",
                   "May 2019 – July 2019", [
                       "Developed a REST API using FastAPI and PostgreSQL to store data from learning management systems",
                       "Explored methods to generate video game dungeons based off of The Legend of Zelda",
                       "Developed a game in Java to test the generated dungeons",
                       "Contributed 50K+ lines of code to an established codebase via Git",
                       "Conducted a human subject study to determine which video game dungeon generation technique is enjoyable"
                   ])
    #document.add_paragraph()
    
    # Add projects section
    add_sectionheader(document, "PROJECTS")
    add_projects(document, "Wildfire Path Predictor", "May 2024 – June 2024", [
        "Leveraged and preprocessed satellite data from MODIS, VIIRS, and IBAND to train an AI model using PyTorch & TorchVision in conjunction with CNN (spatial) and LSTM (temporal) to effectively predict wildfire spread",
        "Implemented a preprocessing pipeline to extract and preprocess hand landmarks using Mediapipe achieving efficient representation of sign language gestures for input into a neural network model"
    ])
    add_projects(document, "GetTrash Garbage Identifier", "May 2024 – May 2024", [
        "Developed a sign language interpreter with machine learning techniques and Scikit-Learn leveraging OpenCV to capture a dataset of over 40,000 images",
        "Implemented a preprocessing pipeline to extract and preprocess hand landmarks using Mediapipe achieving efficient representation of sign language gestures for input into a neural network model"
    ])
    add_projects(document, "Meerkat ASL", "April 2024 – May 2024", [
        "Developed a sign language interpreter with machine learning techniques and Scikit-Learn leveraging OpenCV to capture a dataset of over 40,000 images",
        "Implemented a preprocessing pipeline to extract and preprocess hand landmarks using Mediapipe achieving efficient representation of sign language gestures for input into a neural network model"
    ])
    #document.add_paragraph()
    
    # Add technical skills section
    add_sectionheader(document, "TECHNICAL SKILLS")
    add_technical_skills(document, {
        "Programming Languages": "Python, C++, C, HTML, CSS, JavaScript, Bash, Git, Racket",
        "Technologies": "Flask, OpenCV, Tensorflow, PyTorch, Docker, Kubernetes, Git, Milvus, Rasterio, Streamlit"
    })



    document.save("JakeCV.docx")
    print("Resume created successfully!")

def add_contact_info(document, email, phone, github, linkedin):
    # Add a paragraph with a bottom border to create a horizontal line
    contact_info = document.add_paragraph()
    run = contact_info.add_run(f"{email} ❖| {phone} ❖| {github} ❖| {linkedin}")
    run.font.size = Pt(12)
    run.font.name = 'Garamond'
    contact_info.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    p = contact_info._element
    p_pr = p.get_or_add_pPr()
    p_borders = OxmlElement('w:pBdr')
    p_pr.append(p_borders)
    
    bottom_border = OxmlElement('w:bottom')
    bottom_border.set(qn('w:val'), 'single')
    bottom_border.set(qn('w:sz'), '6')
    bottom_border.set(qn('w:space'), '0')  # Set space before and after to zero
    bottom_border.set(qn('w:color'), '000000')
    p_borders.append(bottom_border)

    # Remove space before and after the paragraph
    contact_info.space_before = Pt(0)
    contact_info.space_after = Pt(0)


def add_sectionheader(document, section_name):
    section = document.add_paragraph()
    run = section.add_run(f"{section_name}")
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = 'Garamond'
    section.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    section.paragraph_format.space_after = Pt(0)  # Remove space after paragraph
    section.paragraph_format.space_before = Pt(12)  # Remove space after paragraph

    p = section._element
    p_pr = p.get_or_add_pPr()
    p_borders = OxmlElement('w:pBdr')
    p_pr.append(p_borders)
    
    bottom_border = OxmlElement('w:bottom')
    bottom_border.set(qn('w:val'), 'single')
    bottom_border.set(qn('w:sz'), '6')
    bottom_border.set(qn('w:space'), '1')
    bottom_border.set(qn('w:color'), '000000')
    p_borders.append(bottom_border)



def add_education(document, university, location, degree, duration, details):
    education_info = document.add_paragraph()
    uni_dur_space = calculate_num_of_space('Garamond Bold', 11, university, duration, 0.5, 0.5)
    run = education_info.add_run(f"{university}{uni_dur_space}{duration}")
    run.bold = True
    run.font.name = 'Garamond'
    run.add_break()
    deg_location_space = calculate_num_of_space('Garamond Italics', 11, degree, location, 0.5, 0.5)
    run = education_info.add_run(f"{degree}{deg_location_space}{location}")
    run.italic = True
    run.font.name = 'Garamond'
    education_info.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    education_info.paragraph_format.space_after = Pt(0)  # Remove space after paragraph
    
    

    for detail in details:
        bullet_paragraph = document.add_paragraph(detail, style='List Bullet')
        for run in bullet_paragraph.runs:
            run.font.name = 'Garamond'
        bullet_paragraph.paragraph_format.space_before = Pt(0)  # Remove space before paragraph
        bullet_paragraph.paragraph_format.space_after = Pt(0)  # Remove space after paragraph
        bullet_paragraph.paragraph_format.line_spacing = Pt(line_spacing)  # Set line spacing
        bullet_paragraph.paragraph_format.left_indent = Inches(0.5)  # Set indentation


    # Turn off "Don't add space between paragraphs of the same style" for this section

def add_experience(document, title, company, location, duration, responsibilities):
    experience_info = document.add_paragraph()
    title_dur_space = calculate_num_of_space('Garamond Bold', 11, title, duration, 0.5, 0.5)
    run = experience_info.add_run(f"{title}{title_dur_space}{duration}")
    run.bold = True
    run.font.name = 'Garamond'
    run.add_break()
    comp_loc_space = calculate_num_of_space('Garamond Italics', 11, company, location, 0.5, 0.5)
    run = experience_info.add_run(f"{company}{comp_loc_space}{location}")
    run.italic = True
    run.font.name = 'Garamond'
    experience_info.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    experience_info.paragraph_format.space_after = Pt(0)  # Remove space after paragraph
    
    for responsibility in responsibilities:
        bullet_paragraph = document.add_paragraph(responsibility, style='List Bullet')
        bullet_paragraph.paragraph_format.space_before = Pt(0)  # Remove space after paragraph
        for run in bullet_paragraph.runs:
            run.font.size = Pt(10)
            run.font.name = 'Garamond'
        bullet_paragraph.paragraph_format.space_after = Pt(para_spacing)  # Remove space after paragraph
        bullet_paragraph.paragraph_format.line_spacing = Pt(line_spacing)  # Set line spacing
        bullet_paragraph.paragraph_format.left_indent = Inches(0.5)  # Set indentation


def add_projects(document, name, duration, description):
    project_info = document.add_paragraph()
    proj_dur_space = calculate_num_of_space('Garamond Bold', 11, name, duration, 0.5, 0.5)
    run = project_info.add_run(f"{name}{proj_dur_space}{duration}")
    run.bold = True
    run.font.name = 'Garamond'
    project_info.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    project_info.paragraph_format.space_after = Pt(0)  # Remove space after paragraph
    

    for desc in description:
        bullet_paragraph = document.add_paragraph(desc, style='List Bullet')
        for run in bullet_paragraph.runs:
            run.font.name = 'Garamond'
            run.font.size = Pt(10)
        bullet_paragraph.paragraph_format.space_after = Pt(para_spacing)  # Remove space after paragraph
        bullet_paragraph.paragraph_format.line_spacing = Pt(line_spacing)  # Set line spacing
        bullet_paragraph.paragraph_format.left_indent = Inches(0.5)  # Set indentation


    # Turn off "Don't add space between paragraphs of the same style" for this section
def add_technical_skills(document, skills):
    skills_info = document.add_paragraph()
    run = skills_info.add_run("Technical Skills")
    run.bold = True
    run.font.name = 'Garamond'
    skills_info.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    skills_info.paragraph_format.space_after = Pt(0)  # Remove space after paragraph

    for category, skill_set in skills.items():
        skill_paragraph = document.add_paragraph(f"{category}: {skill_set}", style='Normal')
        skill_paragraph.paragraph_format.space_before = Pt(0)
        for run in skill_paragraph.runs:
            run.font.name = 'Garamond'
        skill_paragraph.paragraph_format.space_after = Pt(0)  # Remove space after paragraph

create_resume()