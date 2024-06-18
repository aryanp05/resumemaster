from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from CR_spacing import calculate_num_of_space

para_spacing = 6
line_spacing = 14

def add_contact_info(document, contact_info_array):
    # Extract font name and size
    contact_info_list = contact_info_array
    
    # Join the contact info with the separator " ❖| "
    contact_info_str = " | ".join(contact_info_list)

    # Add a paragraph with a bottom border to create a horizontal line
    contact_info = document.add_paragraph()
    run = contact_info.add_run(contact_info_str)
    run.font.size = Pt(12)
    run.font.name = 'Garamond'
    contact_info.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    p = contact_info._element
    p_pr = p.get_or_add_pPr()
    p_borders = OxmlElement('w:pBdr')
    p_pr.append(p_borders)

    bottom_border = OxmlElement('w:bottom')
    bottom_border.set(qn('w:val'), 'single')
    bottom_border.set(qn('w:sz'), '6')
    bottom_border.set(qn('w:space'), '0')
    bottom_border.set(qn('w:color'), '000000')
    p_borders.append(bottom_border)

    # Remove space before and after the paragraph
    contact_info.paragraph_format.space_before = Pt(0)
    contact_info.paragraph_format.space_after = Pt(0)


def add_sectionheader(document, section_name, font, size):
    section = document.add_paragraph()
    run = section.add_run(f"{section_name}")
    run.bold = True
    run.font.size = Pt(size[0])
    run.font.name = font
    section.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    section.paragraph_format.space_after = Pt(0)  # Remove space after paragraph
    section.paragraph_format.space_before = Pt(12)  # xRemove space after paragraph

    p = section._element
    p_pr = p.get_or_add_pPr()
    p_borders = OxmlElement('w:pBdr')
    p_pr.append(p_borders)
    
    bottom_border = OxmlElement('w:bottom')
    bottom_border.set(qn('w:val'), 'single')
    bottom_border.set(qn('w:sz'), '6')
    bottom_border.set(qn('w:space'), '1')
    bottom_border.set(qn('w:color'), '000000')
    p_borders.append(bottom_border)





def add_education(document, university, duration, degree, location, details, font, size):
    education_info = document.add_paragraph()
    uni_dur_space = calculate_num_of_space(font, 'bold', size[1], university, duration, 0.5, 0.5)
    run = education_info.add_run(f"{university}{uni_dur_space}{duration}")
    run.bold = True
    run.font.name = font
    run.add_break()
    deg_location_space = calculate_num_of_space(font, 'italics', size[1], degree, location, 0.5, 0.5)
    run = education_info.add_run(f"{degree}{deg_location_space}{location}")
    run.italic = True
    run.font.size = Pt(size[1])
    run.font.name = font
    education_info.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    education_info.paragraph_format.space_after = Pt(0)  # Remove space after paragraph
    
    

    for detail in details:
        bullet_paragraph = document.add_paragraph(style='List Bullet')
        parts = detail.split('**')
        for i, part in enumerate(parts):
            run = bullet_paragraph.add_run(part)
            run.font.name = font
            run.font.size = Pt(size[2])
            if i % 2 == 1:
                run.bold = True
        bullet_paragraph.paragraph_format.space_before = Pt(0)  # Remove space before paragraph
        bullet_paragraph.paragraph_format.space_after = Pt(0)  # Remove space after paragraph
        bullet_paragraph.paragraph_format.line_spacing = Pt(line_spacing)  # Set line spacing
        bullet_paragraph.paragraph_format.left_indent = Inches(0.5)  # Set indentation



    # Turn off "Don't add space between paragraphs of the same style" for this section

def add_experience(document, title, duration, company, location, responsibilities, font, size):
    experience_info = document.add_paragraph()
    title_dur_space = calculate_num_of_space(font, 'bold', size[1], title, duration, 0.5, 0.5)
    run = experience_info.add_run(f"{title}{title_dur_space}{duration}")
    run.bold = True
    run.font.name = font
    run.add_break()
    comp_loc_space = calculate_num_of_space(font, 'italics', size[1], company, location, 0.5, 0.5)
    run = experience_info.add_run(f"{company}{comp_loc_space}{location}")
    run.italic = True
    run.font.name = font
    experience_info.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    experience_info.paragraph_format.space_after = Pt(0)  # Remove space after paragraph
    
    for responsibility in responsibilities:
        bullet_paragraph = document.add_paragraph(style='List Bullet')
        parts = responsibility.split('**')
        for i, part in enumerate(parts):
            run = bullet_paragraph.add_run(part)
            run.font.name = font
            run.font.size = Pt(size[2])
            if i % 2 == 1:
                run.bold = True
        bullet_paragraph.paragraph_format.space_before = Pt(0)  # Remove space before paragraph
        bullet_paragraph.paragraph_format.space_after = Pt(para_spacing)  # Remove space after paragraph
        bullet_paragraph.paragraph_format.line_spacing = Pt(line_spacing)  # Set line spacing
        bullet_paragraph.paragraph_format.left_indent = Inches(0.5)  # Set indentation



def add_projects(document, name, duration, description, font, size):
    project_info = document.add_paragraph()
    proj_dur_space = calculate_num_of_space(font, 'bold', size[1], name, duration, 0.5, 0.5)
    run = project_info.add_run(f"{name}{proj_dur_space}{duration}")
    run.bold = True
    run.font.name = font
    run.font.size = Pt(size[1])
    project_info.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    project_info.paragraph_format.space_after = Pt(0)  # Remove space after paragraph
    

    for desc in description:
        bullet_paragraph = document.add_paragraph(style='List Bullet')
        parts = desc.split('**')
        for i, part in enumerate(parts):
            run = bullet_paragraph.add_run(part)
            run.font.name = font
            run.font.size = Pt(size[2])
            if i % 2 == 1:
                run.bold = True
        bullet_paragraph.paragraph_format.space_after = Pt(para_spacing)  # Remove space after paragraph
        bullet_paragraph.paragraph_format.line_spacing = Pt(line_spacing)  # Set line spacing
        bullet_paragraph.paragraph_format.left_indent = Inches(0.5)  # Set indentation



    # Turn off "Don't add space between paragraphs of the same style" for this section
def add_technical_skills(document, skills_array, font, size):
    category = skills_array[0]
    skills = skills_array[1]

    # Create a paragraph for the category
    category_paragraph = document.add_paragraph(style='Normal')
    
    # Add category name with bold formatting
    category_run = category_paragraph.add_run(f"{category}: ")
    category_run.bold = True
    category_run.font.name = font
    category_run.font.size = Pt(size[2])

    # Handle skills with potential bold markers
    parts = skills.split('**')
    for i, part in enumerate(parts):
        run = category_paragraph.add_run(part)
        run.font.name = font
        run.font.size = Pt(size[2])
        if i % 2 == 1:
            run.bold = True

    # Adjust formatting for the paragraph
    category_paragraph.paragraph_format.space_before = Pt(0)
    category_paragraph.paragraph_format.space_after = Pt(0)  # Remove space after paragraph